"""
Generate Viva Voce Q&A Preparation Document (.docx)
Based on analysis of the FYP report, rubrics, and project artifacts.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
style.paragraph_format.line_spacing = 1.0
style.paragraph_format.space_after = Pt(2)
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Title Page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('VIVA VOCE Q&A PREPARATION\n')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Coral Reef Health Assessment via\nConvolutional Neural Network-Based Image Analysis')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Muhammad Luqman Haziq bin Mohamad Lofi\n221022249\nBachelor of Computer Engineering\nUniversiti Malaysia Perlis (UniMAP)')
run.font.size = Pt(10.5)

doc.add_page_break()

# Helper functions
def add_section_header(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_question(q_num, question):
    p = doc.add_paragraph()
    run = p.add_run(f'Q{q_num}: {question}')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(153, 0, 0)

def add_bullet(text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    if indent_level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * indent_level)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' {text}')
    else:
        p.add_run(text)

def add_spacer():
    pass

# =====================================================
# SECTION A: ARCHITECTURE & MODEL CHOICE
# =====================================================
add_section_header('SECTION A: ARCHITECTURE & MODEL SELECTION')

# Q1
add_question(1, 'Why did you choose EfficientNetB0 specifically? Why not ResNet50, VGG16, or MobileNet?')
add_bullet('EfficientNetB0 uses compound scaling (scales depth + width + resolution together), so it gets higher accuracy per computation compared to ResNet50 or VGG16.', bold_prefix='Efficiency:')
add_bullet('Zhou et al. (2024) proved that EfficientNetB0-TRVFL beat ResNet50, VGG16, AlexNet, and GoogLeNet on 3 underwater datasets (87.28%, 74.06%, 99.59%).', bold_prefix='Evidence:')
add_bullet('EfficientNetB0 has only ~5.3M parameters. VGG16 has ~138M and ResNet50 has ~25.6M. Our dataset has only 1,582 images, so bigger models would overfit.', bold_prefix='Size matters:')
add_bullet('MobileNet focuses on speed over accuracy. For conservation decisions, classification reliability is more important than speed.', bold_prefix='Why not MobileNet:')
add_bullet('ResNet50 has 5x more parameters than EfficientNetB0. The extra capacity does not bring proportional accuracy gain on small datasets.', bold_prefix='Why not ResNet50:')
add_bullet('Shao et al. (2024): EfficientNet family averaged 82.31% micro F1 vs VGG family 76.53%, while using far fewer parameters (39.75M vs 138.40M).', bold_prefix='More evidence:')

add_spacer()

# Q2
add_question(2, 'Why EfficientNetB0 and not a higher variant like B3, B5, or B7?')
add_bullet('B0 has 5.3M params and uses 224x224 input. This is well-matched to our small dataset of 1,582 images across 3 classes.', bold_prefix='Right size:')
add_bullet('Higher variants have more parameters (B3: 12M, B5: 30M, B7: 66M). More parameters = higher overfitting risk on small data.', bold_prefix='Overfitting risk:')
add_bullet('B0 trains each seed in about 12 minutes on RTX 3070. Higher variants need more VRAM and time, making the 5-seed ensemble impractical.', bold_prefix='Training cost:')
add_bullet('Wang et al. (2024) showed EfficientNet without ensemble scored only 54.25% on coral. The improvement came from our ensemble strategy, not bigger model.', bold_prefix='Key insight:')
add_bullet('Tan & Le (2019) showed B0 already captures the core features. We compensate for limitations through ensembling, not brute-force model scaling.', bold_prefix='Design choice:')

add_spacer()

# Q3
add_question(3, 'Wang et al. (2024) showed EfficientNet scored only 54.25% on coral data. Why did you still use it?')
add_bullet('Wang et al. used a single EfficientNet with no ensemble, no SWA, no oversampling, and no TTA. That is the "vanilla" setup.', bold_prefix='Their setup:')
add_bullet('Our pipeline adds: 5-seed SWA ensemble, 30x Dead class oversampling, cosine decay scheduling, and Multi-Scale TTA.', bold_prefix='Our additions:')
add_bullet('Our baseline single EfficientNetB0 scored 84.91%. After all enhancements, the final ensemble scored 98.11%.', bold_prefix='Our result:')
add_bullet('This validates our research gap: no study combined EfficientNetB0 + multi-seed SWA + multi-scale TTA + Grad-CAM for tri-class coral classification.', bold_prefix='Research gap:')
add_bullet('Using the same base architecture lets us do direct ablation comparison, isolating each pipeline component\'s contribution.', bold_prefix='Methodology benefit:')

add_spacer()

# Q4
add_question(4, 'Why not use Vision Transformers (ViT) or newer architectures like ConvNeXt?')
add_bullet('ViT needs very large datasets (ImageNet-21k scale) to learn spatial relationships. 1,582 images is too small for transformers.', bold_prefix='Data hungry:')
add_bullet('Wang et al. (2024) showed ViT scored lower than ML-Net on the coral bleaching dataset. ViTs struggle with fine-grained texture on small coral data.', bold_prefix='Evidence:')
add_bullet('ConvNeXt modernises CNN design but has higher parameter counts. The marginal accuracy gain does not justify the extra cost for this project.', bold_prefix='ConvNeXt:')
add_bullet('Grad-CAM is natively compatible with CNNs. Applying Grad-CAM to ViTs requires architectural modifications and workarounds.', bold_prefix='Grad-CAM fit:')
add_bullet('EfficientNetB0 keeps sufficient spatial resolution in its final conv layers for meaningful Grad-CAM heatmaps. Confirmed by Shukla et al. (2020) and Korkmaz et al. (2025).', bold_prefix='Explainability:')

add_spacer()

# Q5
add_question(5, 'Can you explain compound scaling in EfficientNet? What makes it different?')
add_bullet('Traditional scaling: increase only depth (ResNet), or only width (WideResNet), or only resolution. One dimension at a time.', bold_prefix='Old approach:')
add_bullet('EfficientNet scales all three together using a compound coefficient. This distributes model capacity across all dimensions evenly.', bold_prefix='New approach:')
add_bullet('Result: EfficientNetB0 gets 77.1% top-1 on ImageNet with only 5.3M params. ResNet50 needs 25.6M for 76.0%.', bold_prefix='Numbers:')
add_bullet('Balanced scaling preserves spatial resolution in later layers. This is critical for Grad-CAM to produce meaningful heatmaps.', bold_prefix='Why it matters:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION B: TRAINING STRATEGY & TECHNIQUES
# =====================================================
add_section_header('SECTION B: TRAINING STRATEGY & TECHNIQUES')

# Q6
add_question(6, 'What is Stochastic Weight Averaging (SWA) and why did you use it?')
add_bullet('SWA averages model weights collected during later training epochs (epochs 26-30). This produces weights that sit in a flatter region of the loss landscape.', bold_prefix='What it does:')
add_bullet('Flatter minima generalise better than sharp minima. A model in a sharp minimum is sensitive to small data changes.', bold_prefix='Why flatter is better:')
add_bullet('We combined SWA with 5 independent seeds (42-46). Each seed converges differently, then SWA smooths each one.', bold_prefix='Our setup:')
add_bullet('Pham-Ngoc et al. (2025) showed a single EfficientNetB0 dropped from F1=89.19% to F1=68% on external data. SWA ensemble directly addresses this gap.', bold_prefix='Literature support:')
add_bullet('Result: baseline single model = 84.91%. Our 5-seed SWA ensemble = 98.11%.', bold_prefix='Our result:')

add_spacer()

# Q7
add_question(7, 'Why 5 seeds specifically? Why not 3, 7, or 10?')
add_bullet('We tested 1, 2, 3, and 5 seeds in our ablation study.', bold_prefix='Tested:')
add_bullet('1 seed: 96.23% accuracy, 6 errors. 3 seeds: 98.11%, 3 errors. 5 seeds: 98.11%, 3 errors.', bold_prefix='Results:')
add_bullet('Performance plateaus at 3 seeds. No accuracy gain between 3 and 5 seeds.', bold_prefix='Plateau:')
add_bullet('We keep 5 seeds because broader weight-space averaging gives more stable predictions, even without accuracy gain.', bold_prefix='Why keep 5:')
add_bullet('More than 5 gives diminishing returns. Each seed takes ~12 min. 10 seeds = ~120 min with no expected benefit.', bold_prefix='Why not more:')
add_bullet('Total training: 5 seeds x 12 min = ~60 min. Feasible on RTX 3070.', bold_prefix='Practical:')

add_spacer()

# Q8
add_question(8, 'What is Hard-Example Oversampling and why did you use it instead of SMOTE or class weighting alone?')
add_bullet('Problem: Healthy=712, Bleached=720, Dead=150. Dead class is only 9.5% of the data.', bold_prefix='Imbalance:')
add_bullet('Our approach: duplicate the hardest Dead samples at 30x rate. These are samples the model gets wrong during early training.', bold_prefix='What we did:')
add_bullet('SMOTE creates synthetic samples by interpolating in feature space. For images, this can create non-meaningful fake images.', bold_prefix='Why not SMOTE:')
add_bullet('Class weighting alone adjusts the loss function but treats all Dead samples equally. It does not focus on the hardest cases.', bold_prefix='Why not class weighting alone:')
add_bullet('Hard-Example Oversampling forces the model to repeatedly learn from the most confusing Dead samples. This directly improves the Bleached-Dead boundary.', bold_prefix='Advantage:')
add_bullet('Result: Dead class precision = 1.000 (zero false positives). Dead recall = 0.9333 (14/15 correct).', bold_prefix='Result:')

add_spacer()

# Q9
add_question(9, 'Why use cosine decay learning rate schedule instead of step decay or constant LR?')
add_bullet('Cosine decay smoothly reduces the learning rate following a cosine curve from start to near-zero.', bold_prefix='What it does:')
add_bullet('Step decay has abrupt drops. These sudden changes can destabilise training with large gradient jumps.', bold_prefix='Why not step decay:')
add_bullet('Constant LR can overshoot the optimal weights in later epochs when fine-tuning needs small adjustments.', bold_prefix='Why not constant:')
add_bullet('Cosine decay gives aggressive learning early (when the model needs big updates) and fine-grained refinement later (when small adjustments matter).', bold_prefix='Best of both:')
add_bullet('Our validation loss decreased smoothly over 30 epochs with no upward spike. This confirms effective convergence.', bold_prefix='Our evidence:')

add_spacer()

# Q10
add_question(10, 'Explain transfer learning. Why not train from scratch?')
add_bullet('Transfer learning uses a model pretrained on ImageNet (14M+ images, 1000 classes) as a starting point, then fine-tunes on our coral data.', bold_prefix='Definition:')
add_bullet('ImageNet pretraining already learns universal visual features like edges, textures, and colour gradients. These are directly useful for coral texture and colour discrimination.', bold_prefix='Why it works:')
add_bullet('Training from scratch on 1,582 images would severely overfit. Not enough data to learn robust features from zero.', bold_prefix='Why not from scratch:')
add_bullet('Pan & Yang (2010): transfer learning is the standard approach for small datasets in environmental image analysis.', bold_prefix='Literature:')
add_bullet('Our strategy: freeze the backbone first, then unfreeze the top 100 layers, then fine-tune with a low learning rate to adapt to coral features.', bold_prefix='Our approach:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION C: INFERENCE & TTA
# =====================================================
add_section_header('SECTION C: INFERENCE STRATEGY & TEST-TIME AUGMENTATION')

# Q11
add_question(11, 'What is Multi-Scale Test-Time Augmentation (TTA) and how does it work?')
add_bullet('At test time, each image is processed at 2 resolutions: 224x224 (native) and 256x256 (upscaled). Each also gets a horizontal flip. That makes 4 views total.', bold_prefix='How it works:')
add_bullet('All 4 views pass through all 5 ensemble models. That gives 20 predictions per image.', bold_prefix='Total predictions:')
add_bullet('The 20 predictions are averaged (probability averaging) to produce the final class prediction.', bold_prefix='Aggregation:')
add_bullet('This ensures the model is robust to different scales and orientations. Underwater coral images vary in camera distance and angle.', bold_prefix='Purpose:')
add_bullet('Kandel & Castelli (2021): TTA consistently improved CNN performance across 50 experiments. Average voting is the most reliable aggregation.', bold_prefix='Literature:')
add_bullet('Oza et al. (2024): EfficientNet-B7 with TTA got F1=0.9955 vs F1=0.9934 without TTA.', bold_prefix='Evidence:')

add_spacer()

# Q12
add_question(12, 'Your ablation shows TTA did not improve accuracy at 5 seeds. Why keep it?')
add_bullet('TTA was neutral at 5 seeds (98.11% with and without). It did not hurt anything.', bold_prefix='Accuracy:')
add_bullet('TTA adds robustness against orientation, scale, and lighting variations typical in underwater coral images.', bold_prefix='Robustness:')
add_bullet('At lower seed counts (1-2 seeds), TTA was actually unstable. At 5 seeds, the ensemble is stable enough to absorb TTA variance.', bold_prefix='Stability:')
add_bullet('Removing TTA would sacrifice a proven robustness mechanism for zero accuracy cost. Better to keep it as a safety margin for deployment.', bold_prefix='Engineering choice:')

add_spacer()

# Q13
add_question(13, 'Why only 2 scales (224 and 256)? Why not more?')
add_bullet('224x224 is EfficientNetB0\'s native input. 256x256 gives slightly more spatial context (centre-cropped to 224).', bold_prefix='Rationale:')
add_bullet('Larger scales like 384 or 512 would need downsampling back to 224, introducing interpolation artefacts that could damage texture features.', bold_prefix='Why not larger:')
add_bullet('More scales = more forward passes = slower inference. Each extra scale adds latency per image.', bold_prefix='Speed cost:')
add_bullet('Our 2-scale + flip gives 4 views per model. That is enough diversity for reliable averaging.', bold_prefix='Sufficient:')
add_bullet('Mean inference time: 10.38 ms/image. Fast enough for real-time web app use.', bold_prefix='Speed result:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION D: DATASET & PREPROCESSING
# =====================================================
add_section_header('SECTION D: DATASET & DATA PREPROCESSING')

# Q14
add_question(14, 'Why use the BHD Coral Dataset? Are there better datasets available?')
add_bullet('BHD Coral Dataset (Jamil, 2021): 1,582 images, 3 classes (Healthy=712, Bleached=720, Dead=150), publicly available on Kaggle.', bold_prefix='Dataset:')
add_bullet('It provides a clear tri-class taxonomy matching the biological progression: healthy to bleached to dead.', bold_prefix='Good fit:')
add_bullet('Publicly available and reproducible. Any researcher can download it and replicate our work.', bold_prefix='Reproducibility:')
add_bullet('Shao et al. (2024) used 20,000+ images across 8 classes. That is a different, more complex problem beyond our tri-class scope.', bold_prefix='Alternatives:')
add_bullet('Limitation acknowledged: the dataset comes from limited geographic regions. Model generalisation to other reefs is not validated.', bold_prefix='Limitation:')

add_spacer()

# Q15
add_question(15, 'How did you handle the severe class imbalance? Dead class has only 150 images vs 712/720.')
add_bullet('Dead class is only 9.5% of total data. Without intervention, the model would bias toward Healthy/Bleached.', bold_prefix='Problem:')
add_bullet('Strategy 1: Hard-Example Oversampling at 30x for Dead class. Duplicate the hardest Dead samples.', bold_prefix='Main fix:')
add_bullet('Strategy 2: Stratified splitting preserves class proportions across train/val/test (80:10:10).', bold_prefix='Split:')
add_bullet('Applied only to the training set. Validation and test sets stay unmodified. No data leakage.', bold_prefix='No leakage:')
add_bullet('Result: Dead precision = 1.000, recall = 0.9333, F1 = 0.9655. Comparable to majority classes despite 5x fewer test samples.', bold_prefix='Result:')
add_bullet('In real use: missing Dead coral (low recall) delays conservation. False Dead alarms (low precision) waste resources. We optimise for both.', bold_prefix='Real impact:')

add_spacer()

# Q16
add_question(16, 'Why 80:10:10 split? Why not 70:15:15 or k-fold cross-validation?')
add_bullet('80% training maximises learning data for a small dataset. 10% validation for tuning. 10% test (159 images) for final evaluation.', bold_prefix='Rationale:')
add_bullet('With only 1,582 images and Dead at 150, giving more to val/test would further shrink the training set.', bold_prefix='Dataset constraint:')
add_bullet('15% test would give about 23 Dead test samples vs current 15. Marginal improvement does not justify reduced training data.', bold_prefix='Why not 70:15:15:')
add_bullet('K-fold is too expensive with 5-seed ensemble. 5 folds x 5 seeds = 25 full training runs.', bold_prefix='Why not k-fold:')
add_bullet('Fixed split with seed=42 ensures exact reproducibility. Anyone can replicate the same partition.', bold_prefix='Reproducible:')

add_spacer()

# Q17
add_question(17, 'What augmentation techniques did you use and why those specifically?')
add_bullet('Random rotation (+-20 degrees), horizontal flip, random zoom (90-110%). Applied only during training.', bold_prefix='Techniques:')
add_bullet('Rotation: simulates different diver/ROV camera angles underwater.', bold_prefix='Why rotation:')
add_bullet('Horizontal flip: corals have no left-right orientation. Doubles effective training data at zero label cost.', bold_prefix='Why flip:')
add_bullet('Zoom: simulates varying camera-to-coral distances typical in underwater photography.', bold_prefix='Why zoom:')
add_bullet('No colour jittering: underwater images already vary in colour due to depth. Extra colour changes could remove important cues like bleaching whiteness.', bold_prefix='Why no colour aug:')
add_bullet('No vertical flip: underwater images have natural gravity orientation. Vertical flip creates unrealistic upside-down coral images.', bold_prefix='Why no vertical flip:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION E: EXPLAINABILITY (GRAD-CAM)
# =====================================================
add_section_header('SECTION E: EXPLAINABILITY & GRAD-CAM')

# Q18
add_question(18, 'Why Grad-CAM specifically? Why not LIME, SHAP, or other XAI methods?')
add_bullet('Grad-CAM generates class-specific heatmaps using gradients from the last conv layer. It works natively with any CNN.', bold_prefix='What it does:')
add_bullet('LIME needs to perturb the image and run many forward passes. Too slow for real-time web deployment.', bold_prefix='Why not LIME:')
add_bullet('SHAP needs even more inference passes per image. Not practical for per-image real-time explanations.', bold_prefix='Why not SHAP:')
add_bullet('Grad-CAM needs only a single backward pass. Fast enough for real-time inference in our Flask app (10.38 ms/image).', bold_prefix='Speed:')
add_bullet('Grad-CAM + EfficientNet validated by: Shukla et al. (2020), Korkmaz et al. (2025), Montalbo & Alon (2021), Alqhatani et al. (2025).', bold_prefix='Literature:')
add_bullet('Singh et al. (2022) showed Grad-CAM works with TTA-based pipelines. Directly applicable to our setup.', bold_prefix='TTA compatible:')

add_spacer()

# Q19
add_question(19, 'How do your Grad-CAM results prove the model is looking at the right features?')
add_bullet('Healthy: activation on pigmented branching/fan coral. This is the zooxanthellae-rich tissue that gives healthy coral its colour.', bold_prefix='Healthy:')
add_bullet('Bleached: broad activation across pale coral surface. Matches the spatially diffuse nature of thermal bleaching (whole-colony pigment loss).', bold_prefix='Bleached:')
add_bullet('Dead: focused activation on exposed skeletal formations. The model relies on structural texture since there is no living tissue colour.', bold_prefix='Dead:')
add_bullet('Baseline comparison: the single model focused on a clownfish for Healthy classification (75.2% confidence). It used the wrong feature entirely.', bold_prefix='Baseline failure:')
add_bullet('Ensemble confidence: 96.6%-99.4% vs baseline: 45.5%-75.2%. Higher confidence correlates with more focused, biologically correct attention.', bold_prefix='Confidence gap:')

add_spacer()

# Q20
add_question(20, 'The baseline model classified a Healthy coral correctly but Grad-CAM showed it focused on a clownfish. What does this mean?')
add_bullet('The model got the right answer for the wrong reason. It learned a shortcut: clownfish appear with healthy coral in training data.', bold_prefix='Wrong reason:')
add_bullet('This is why accuracy alone is not enough. A model can be "right for the wrong reasons."', bold_prefix='Accuracy trap:')
add_bullet('In deployment: if a healthy coral image has no fish, the baseline would likely fail. The shortcut does not generalise.', bold_prefix='Deployment risk:')
add_bullet('This directly supports Objective 2: "To apply Grad-CAM to support interpretation and validation."', bold_prefix='Objective 2:')
add_bullet('Borjali et al. (2020) warned that DL models can attend to irrelevant background instead of meaningful biological indicators.', bold_prefix='Literature:')
add_bullet('Our ensemble Grad-CAM shows activation on coral tissue, not fish. The ensemble learned genuine coral features.', bold_prefix='Ensemble fix:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION F: RESULTS & EVALUATION
# =====================================================
add_section_header('SECTION F: RESULTS & PERFORMANCE EVALUATION')

# Q21
add_question(21, 'Explain your final results. What do the numbers mean?')
add_bullet('Test accuracy: 98.11%. That is 156 out of 159 test images correctly classified.', bold_prefix='Accuracy:')
add_bullet('Only 3 errors: 2 Bleached predicted as Healthy, 1 Dead predicted as Bleached.', bold_prefix='Errors:')
add_bullet('Macro F1: 0.9769. This treats all 3 classes equally regardless of sample count. Proves balanced performance.', bold_prefix='Macro F1:')
add_bullet('Per-class F1: Healthy=0.9863, Bleached=0.9790, Dead=0.9655. All above 96.5%.', bold_prefix='Class breakdown:')
add_bullet('Dead precision = 1.000. Zero false Dead predictions. No non-Dead sample was ever wrongly called Dead.', bold_prefix='Dead precision:')
add_bullet('Improvement: +13.20 percentage points over baseline (84.91% to 98.11%).', bold_prefix='Improvement:')

add_spacer()

# Q22
add_question(22, 'Why use Macro F1-score instead of just accuracy?')
add_bullet('Accuracy is misleading with imbalanced data. A model predicting only Healthy/Bleached could get ~90.5% accuracy while completely failing on Dead.', bold_prefix='Accuracy problem:')
add_bullet('Macro F1 computes F1 per class, then averages. Gives equal weight to Dead (15 samples) and Healthy (72 samples).', bold_prefix='What Macro F1 does:')
add_bullet('You cannot get high Macro F1 if any class is neglected. It penalises poor minority class performance.', bold_prefix='Why it matters:')
add_bullet('F1 = harmonic mean of precision and recall. Penalises extreme differences between the two.', bold_prefix='F1 definition:')
add_bullet('We also report weighted F1 (0.9810) for completeness. But Macro F1 is our primary metric.', bold_prefix='Both reported:')

add_spacer()

# Q23
add_question(23, 'You have only 15 Dead coral test samples. Is this statistically reliable?')
add_bullet('Honest answer: each Dead error shifts recall by ~6.67 percentage points (1/15). Much more than Healthy/Bleached at ~1.39 (1/72).', bold_prefix='Limitation:')
add_bullet('Dead recall of 93.33% (14/15) carries wider statistical uncertainty than majority class results.', bold_prefix='Uncertainty:')
add_bullet('But Dead precision = 1.000 (zero false positives). No sample was ever wrongly predicted as Dead.', bold_prefix='Counter point:')
add_bullet('The limitation comes from the source dataset (only 150 Dead images total), not from our methodology.', bold_prefix='Root cause:')
add_bullet('Multiple evidence sources converge: ablation study, 5-seed consistency, Grad-CAM validation all support the result.', bold_prefix='Converging evidence:')
add_bullet('Future work: expand Dead data collection and cross-regional validation for broader confidence.', bold_prefix='Future fix:')

add_spacer()

# Q24
add_question(24, 'Explain the 3 misclassifications. What caused them?')
add_bullet('Error 1 & 2: Two Bleached samples predicted as Healthy. Grad-CAM showed activation displaced toward peripheral edges and reef fauna, not coral.', bold_prefix='Bleached as Healthy:')
add_bullet('Root cause: scene-level interference. Reef fauna co-occurring with coral redirected the model\'s attention.', bold_prefix='Cause:')
add_bullet('Error 3: One Dead sample predicted as Bleached at 51.3% confidence. Grad-CAM showed correct localisation but wrong class.', bold_prefix='Dead as Bleached:')
add_bullet('Root cause: genuine visual ambiguity. Bleached and dead coral share reduced pigmentation and structural degradation.', bold_prefix='Cause:')
add_bullet('All 3 errors had sub-55% confidence. The ensemble flagged its own uncertainty. Supports confidence thresholding in deployment.', bold_prefix='Low confidence:')
add_bullet('No sample was ever falsely predicted as Dead. Zero false Dead positives.', bold_prefix='Important:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION G: ABLATION STUDY
# =====================================================
add_section_header('SECTION G: ABLATION STUDY')

# Q25
add_question(25, 'Explain your ablation study. What did it prove?')
add_bullet('Tested 8 configurations: 1, 2, 3, 5 seeds, each with and without Multi-Scale TTA.', bold_prefix='Design:')
add_bullet('Main finding: ensemble size is the dominant performance driver, not TTA.', bold_prefix='Key finding:')
add_bullet('1 seed: 96.23%, 6 errors. 3 seeds: 98.11%, 3 errors. 5 seeds: 98.11%, 3 errors.', bold_prefix='Progression:')
add_bullet('Performance plateaus at 3 seeds. 5 seeds kept for stability, not further accuracy.', bold_prefix='Plateau:')
add_bullet('TTA was unstable at low seed counts. 2-seed + TTA gave worst result (6 errors). TTA only stabilises at 5 seeds.', bold_prefix='TTA instability:')

add_spacer()

# Q26
add_question(26, 'Why was TTA harmful at 2 seeds but neutral at 5 seeds?')
add_bullet('2-seed ensemble has limited model diversity. Only 2 weight solutions are averaged together.', bold_prefix='Root cause:')
add_bullet('TTA introduces geometric input variance (scale changes + flip). With only 2 models, this extra variance cannot be reliably cancelled out.', bold_prefix='Mechanism:')
add_bullet('At 2 seeds + TTA: the extra uncertainty from transformed inputs exceeds the averaging capacity. Result: more errors (6 vs 5 without TTA).', bold_prefix='2-seed result:')
add_bullet('At 5 seeds: the prediction surface is stable enough to absorb TTA variance. TTA contributes robustness without causing instability.', bold_prefix='5-seed result:')
add_bullet('Think of it like signal averaging: more independent signals (seeds) are needed to reliably extract signal from noise (TTA variance).', bold_prefix='Analogy:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION H: DEPLOYMENT & WEB APPLICATION
# =====================================================
add_section_header('SECTION H: DEPLOYMENT & WEB APPLICATION')

# Q27
add_question(27, 'Explain your web application deployment. Why Flask?')
add_bullet('Flask is a lightweight Python web framework. It integrates seamlessly with TensorFlow/Keras without needing a different language.', bold_prefix='Why Flask:')
add_bullet('Our entire stack is Python: training, evaluation, Grad-CAM, and deployment. One language for everything.', bold_prefix='Consistency:')
add_bullet('Assaduzzaman et al. (2025) also deployed their EfficientNet model via Flask for tomato disease classification. Validated approach.', bold_prefix='Literature:')
add_bullet('Features: drag-and-drop image upload, real-time classification, confidence scores, and Grad-CAM overlay display.', bold_prefix='Features:')
add_bullet('Mean inference time: 10.38 ms per image. Well within real-time requirements.', bold_prefix='Speed:')
add_bullet('Limitation: runs on a single local workstation. Not scalable for multi-user or cloud deployment.', bold_prefix='Limitation:')

add_spacer()

# Q28
add_question(28, 'How does the web application integrate Grad-CAM in real time?')
add_bullet('User uploads coral image. Flask backend preprocesses it (resize to 224x224, normalise).', bold_prefix='Step 1:')
add_bullet('Image passes through all 5 ensemble models. Grad-CAM is computed via a single backward pass on the final conv layer.', bold_prefix='Step 2:')
add_bullet('Heatmap is upsampled to 224x224 using bilinear interpolation. JET colourmap overlay is applied via OpenCV.', bold_prefix='Step 3:')
add_bullet('User sees: predicted class label, confidence percentage, and Grad-CAM heatmap overlay. Full transparency of the model\'s decision.', bold_prefix='Output:')
add_bullet('Total latency: about 10 ms for inference plus minimal overhead for Grad-CAM and overlay. Real-time experience.', bold_prefix='Speed:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION I: RESEARCH GAP & CONTRIBUTIONS
# =====================================================
add_section_header('SECTION I: RESEARCH GAP & CONTRIBUTIONS')

# Q29
add_question(29, 'What is your research gap? How does your work fill it?')
add_bullet('No existing study combined EfficientNetB0 + multi-seed SWA + multi-scale TTA + Grad-CAM for tri-class coral health classification.', bold_prefix='Gap:')
add_bullet('Zhou et al. (2024): used EfficientNetB0 for underwater tasks but without ensemble or explainability.', bold_prefix='Gap evidence 1:')
add_bullet('Shao et al. (2024): used ensemble for coral but with Swin Transformer + EfficientNet-B7, not B0 with SWA.', bold_prefix='Gap evidence 2:')
add_bullet('Wang et al. (2024): showed vanilla EfficientNet fails on coral (54.25%) but did not propose ensemble/SWA solutions.', bold_prefix='Gap evidence 3:')
add_bullet('Kallipolitis et al. (2021): EfficientNet ensembles with Grad-CAM but for histopathology, not coral classification.', bold_prefix='Gap evidence 4:')
add_bullet('Our contribution: first study integrating all four components into one pipeline specifically for coral health assessment.', bold_prefix='Our contribution:')

add_spacer()

# Q30
add_question(30, 'What are the limitations of your study?')
add_bullet('BHD dataset from limited reef regions. Model generalisation to other oceans is not validated.', bold_prefix='Geography:')
add_bullet('Only 15 Dead test samples. Each error shifts metrics by ~6.67 percentage points.', bold_prefix='Statistical:')
add_bullet('Static images only. No video or temporal analysis of reef degradation trends.', bold_prefix='No temporal:')
add_bullet('Flask app runs on single workstation. Not scalable for multi-user or field deployment.', bold_prefix='Deployment:')
add_bullet('Trained and tested on similar camera systems. Performance under different imaging conditions is unknown.', bold_prefix='Device:')
add_bullet('Only Grad-CAM used. Multi-method XAI (LIME, SHAP, Grad-CAM++) could give more comprehensive interpretability.', bold_prefix='XAI scope:')

add_spacer()

# Q31
add_question(31, 'What are your recommendations for future work?')
add_bullet('Knowledge distillation: compress 5-model ensemble into single lightweight model for edge devices (NVIDIA Jetson, ARM).', bold_prefix='Compression:')
add_bullet('Active learning: let the model flag low-confidence predictions for expert annotation. Improves the Bleached-Dead boundary over time.', bold_prefix='Active learning:')
add_bullet('Domain adaptation: use adversarial feature alignment and style transfer to generalise across different reef regions.', bold_prefix='Generalisation:')
add_bullet('AUV integration: embed the model in Autonomous Underwater Vehicles for field-portable reef assessment without divers.', bold_prefix='Field deployment:')
add_bullet('Federated learning: multiple marine institutions train locally without sharing raw data. Collaborative improvement with data privacy.', bold_prefix='Federated:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION J: PROBLEM STATEMENT & OBJECTIVES
# =====================================================
add_section_header('SECTION J: PROBLEM STATEMENT, OBJECTIVES & SCOPE')

# Q32
add_question(32, 'What are your three problem statements?')
add_bullet('PS1: Manual analysis cannot handle the volume of underwater imagery from modern ROVs and AUVs. (Gonzalez-Rivero et al., 2020; Mahmood et al., 2016)', bold_prefix='Volume:')
add_bullet('PS2: Deep learning models are black boxes. They may attend to background noise instead of coral features. No visibility into decisions. (Borjali et al., 2020; Selvaraju et al., 2020)', bold_prefix='Interpretability:')
add_bullet('PS3: Using only accuracy is misleading for imbalanced data. Dead coral can be systematically missed while accuracy looks high. (Borjali et al., 2020)', bold_prefix='Evaluation:')

add_spacer()

# Q33
add_question(33, 'How does each objective map to your methodology and results?')
add_bullet('Objective 1 (develop model): EfficientNetB0 + 5-seed SWA ensemble + oversampling. Result: 98.11% accuracy, 3 errors.', bold_prefix='Obj 1:')
add_bullet('Objective 2 (visual explanation): Grad-CAM on final conv layer. Result: biologically coherent heatmaps. Baseline comparison proved ensemble learns correct features.', bold_prefix='Obj 2:')
add_bullet('Objective 3 (multi-metric evaluation): Accuracy + Precision + Recall + F1 + Macro F1 + Confusion Matrix. Result: Macro F1 = 0.9769, all class F1 > 0.965.', bold_prefix='Obj 3:')

add_spacer()

# Q34
add_question(34, 'How did your project scope change from FYP1 to FYP2?')
add_bullet('Model: generic CNN (unspecified) became EfficientNetB0 with ImageNet pretraining.', bold_prefix='Model:')
add_bullet('Training: single model, single seed became Hard-Example Oversampling + 5-seed SWA ensemble.', bold_prefix='Training:')
add_bullet('Inference: standard forward pass became Multi-Scale TTA (224+256, original+flip).', bold_prefix='Inference:')
add_bullet('Explainability: "visual explanation method" (unspecified) became Grad-CAM.', bold_prefix='XAI:')
add_bullet('Deployment: not in original scope. Added Flask web app with live inference and Grad-CAM overlay.', bold_prefix='Deployment:')
add_bullet('Added Objective 4: web-based deployment (beyond the original 3 objectives).', bold_prefix='New objective:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION K: TECHNICAL DEPTH (CHALLENGING QUESTIONS)
# =====================================================
add_section_header('SECTION K: CHALLENGING TECHNICAL QUESTIONS')

# Q35
add_question(35, 'How do you know your model is not overfitting?')
add_bullet('Validation loss stays below training loss for all 30 epochs. No upward spike.', bold_prefix='Loss curves:')
add_bullet('Training loss > validation loss is expected because Hard-Example Oversampling makes training batches harder than the clean validation set.', bold_prefix='Explanation:')
add_bullet('Accuracy gap stabilised within +/-0.40 percentage points by epoch 22. No widening trend.', bold_prefix='Accuracy:')
add_bullet('Test accuracy 98.11% on held-out data confirms generalisation. An overfitted model would show high train accuracy but low test accuracy.', bold_prefix='Test proof:')
add_bullet('SWA itself acts as regularisation. Averaging weights across multiple training trajectories reduces sensitivity to any single path.', bold_prefix='SWA regularises:')
add_bullet('Data augmentation (rotation, flip, zoom) also acts as implicit regularisation during training.', bold_prefix='Augmentation:')

add_spacer()

# Q36
add_question(36, 'What is the confusion matrix telling us? Interpret it.')
add_bullet('Healthy: 72/72 correct (100% recall). Zero missed. But 2 Bleached samples wrongly predicted as Healthy.', bold_prefix='Healthy row:')
add_bullet('Bleached: 70/72 correct (97.2% recall). 2 predicted as Healthy due to visual ambiguity at the Healthy-Bleached boundary.', bold_prefix='Bleached row:')
add_bullet('Dead: 14/15 correct (93.3% recall). 1 predicted as Bleached due to boundary confusion between visually similar degraded states.', bold_prefix='Dead row:')
add_bullet('Dead column: no sample was predicted as Dead incorrectly. Dead precision = 1.000.', bold_prefix='Dead column:')
add_bullet('Errors only flow along adjacent health states (Dead to Bleached, Bleached to Healthy). No "skip" errors like Dead to Healthy.', bold_prefix='Error direction:')
add_bullet('This directional pattern aligns with biology. Adjacent states share more visual features than distant states.', bold_prefix='Biology match:')

add_spacer()

# Q37
add_question(37, 'Your validation loss is below training loss throughout training. Is this normal?')
add_bullet('Yes. This is the expected result of our high-regularisation training strategy.', bold_prefix='Answer:')
add_bullet('Cause 1: Hard-Example Oversampling at 30x makes training batches artificially harder.', bold_prefix='Oversampling:')
add_bullet('Cause 2: Data augmentation adds noise to training images only. Each batch is harder than clean validation data.', bold_prefix='Augmentation:')
add_bullet('Cause 3: Dropout 0.4 disables 40% of neurons during training. At validation, all neurons are active (full capacity).', bold_prefix='Dropout:')
add_bullet('Validation sees clean images with full network capacity. Naturally easier.', bold_prefix='Why validation is lower:')
add_bullet('If this were underfitting, both training and validation would perform poorly. But validation accuracy exceeds 98%.', bold_prefix='Not underfitting:')

add_spacer()

# Q38
add_question(38, 'How does your ensemble averaging work mathematically?')
add_bullet('Each of the 5 seeds produces a softmax probability vector [P_healthy, P_bleached, P_dead] per image.', bold_prefix='Step 1:')
add_bullet('With TTA: each seed processes 4 views (2 scales x 2 orientations) = 20 probability vectors per image total.', bold_prefix='Step 2:')
add_bullet('All 20 vectors are element-wise averaged: final_P = (1/20) x sum of all P_i for each class.', bold_prefix='Step 3:')
add_bullet('Final prediction = argmax(final_P). The class with highest averaged probability wins.', bold_prefix='Step 4:')
add_bullet('Probability averaging (soft voting) is better than majority voting (hard voting) because it preserves confidence information.', bold_prefix='Why soft voting:')
add_bullet('Lakshminarayanan et al. (2017): deep ensembles with probability averaging give well-calibrated uncertainty estimates.', bold_prefix='Literature:')

add_spacer()

# Q39
add_question(39, 'What hardware did you use? Could this run on cheaper hardware?')
add_bullet('Training: NVIDIA RTX 3070 GPU (8 GB VRAM), CUDA 11.8, cuDNN 8.6.', bold_prefix='Hardware:')
add_bullet('Each seed: about 12 min training x 5 seeds = about 60 min total.', bold_prefix='Time:')
add_bullet('EfficientNetB0 is lightweight (~5.3M params). Could train on GPUs with 4 GB VRAM like GTX 1650 or RTX 2060.', bold_prefix='Cheaper options:')
add_bullet('Inference: 10.38 ms per image. Could run on CPU for non-real-time use, though GPU is recommended for the web app.', bold_prefix='Inference:')
add_bullet('Future work: knowledge distillation for edge deployment on NVIDIA Jetson or ARM processors.', bold_prefix='Future:')

add_spacer()

# Q40
add_question(40, 'Why did you focus on only 3 classes? Real coral reefs have many conditions.')
add_bullet('Healthy, Bleached, Dead represent the fundamental biological progression of coral thermal stress (Hoegh-Guldberg et al., 2007).', bold_prefix='Biology:')
add_bullet('These 3 states are the primary categories used in conservation and reef management decisions.', bold_prefix='Conservation:')
add_bullet('The BHD dataset is specifically labelled for these 3 classes with consistent and reliable annotations.', bold_prefix='Data available:')
add_bullet('More classes (partially bleached, diseased, recovering) would require expert annotation and much larger datasets.', bold_prefix='Scope constraint:')
add_bullet('Shao et al. (2024) used 8 classes but needed 20,000+ images. More classes require significantly more data.', bold_prefix='Evidence:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION L: METHODOLOGY PIPELINE & FLOWCHART
# =====================================================
add_section_header('SECTION L: METHODOLOGY PIPELINE & FLOWCHART')

# Q41
add_question(41, 'Walk me through your complete methodology pipeline from start to finish.')
add_bullet('Stage 1: Data Acquisition. Get BHD Coral Dataset from Kaggle (1,582 images, 3 classes).', bold_prefix='Step 1:')
add_bullet('Stage 2: Preprocessing. BGR to RGB, resize to 224x224, normalise to [0,1], one-hot encode labels.', bold_prefix='Step 2:')
add_bullet('Stage 3: Stratified splitting. 80:10:10 (train:val:test) with seed=42, preserving class proportions.', bold_prefix='Step 3:')
add_bullet('Stage 4: Augmentation. Random rotation, horizontal flip, random zoom on training set only.', bold_prefix='Step 4:')
add_bullet('Stage 5: Hard-Example Oversampling. 30x for Dead class hard samples in training set.', bold_prefix='Step 5:')
add_bullet('Stage 6: Model development. EfficientNetB0 backbone (ImageNet) + custom head (GAP, Dense 256, Dropout, Dense 3 + softmax).', bold_prefix='Step 6:')
add_bullet('Stage 7: Training. 5 seeds (42-46), 30 epochs each, cosine decay LR, SWA weight collection at epochs 26-30.', bold_prefix='Step 7:')
add_bullet('Stage 8: Hyperparameter tuning. Iterative validation against accuracy/loss/F1 criteria.', bold_prefix='Step 8:')
add_bullet('Stage 9: Evaluation. Multi-Scale TTA at 224 and 256 with flip, then probability averaging across 5 seeds.', bold_prefix='Step 9:')
add_bullet('Stage 10: Output. Classification report + Confusion matrix + Grad-CAM heatmaps + Flask web app.', bold_prefix='Step 10:')

add_spacer()

# Q42
add_question(42, 'What loss function did you use and why?')
add_bullet('Categorical cross-entropy. The standard loss for multi-class classification with one-hot encoded labels.', bold_prefix='Loss function:')
add_bullet('It measures how far the predicted probability distribution is from the true class.', bold_prefix='What it measures:')
add_bullet('Works well with softmax output layer. Gradients flow cleanly back through the network.', bold_prefix='Compatibility:')
add_bullet('Combined with class weighting to give Dead class samples higher loss contribution.', bold_prefix='Class weighting:')
add_bullet('Why not focal loss: focal loss downweights easy examples, but our Hard-Example Oversampling already handles this. Using both could over-correct.', bold_prefix='Why not focal loss:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION M: COMPARISON WITH LITERATURE
# =====================================================
add_section_header('SECTION M: COMPARISON WITH RELATED WORKS')

# Q43
add_question(43, 'How does your 98.11% accuracy compare with other studies in the literature?')
add_bullet('Zhou et al. (2024): EfficientNetB0-TRVFL got 87.28% on MLC2008, 74.06% on MLC2009, 99.59% on Fish-gres. Different underwater tasks, not directly comparable.', bold_prefix='Zhou:')
add_bullet('Shao et al. (2024): ensemble (Swin+EfficientNet-B7) got 86.39% micro F1 on 8-class coral (20,000+ images). More complex task.', bold_prefix='Shao:')
add_bullet('Wang et al. (2024): ML-Net got 86.35% on binary coral bleaching. Our tri-class task is harder but we achieved higher performance.', bold_prefix='Wang:')
add_bullet('Alqhatani et al. (2025): modified EfficientNetB0 for lung cancer staging got 99.24%. Different domain but similar architecture.', bold_prefix='Alqhatani:')
add_bullet('Direct comparison is hard because of different datasets, class counts, and tasks. But 98.11% on tri-class coral is competitive.', bold_prefix='Fair note:')
add_bullet('The baseline comparison (84.91% to 98.11%) on the same dataset is the most scientifically valid comparison.', bold_prefix='Best comparison:')

add_spacer()

# Q44
add_question(44, 'What ethical considerations exist for deploying AI in ecological monitoring?')
add_bullet('False Dead predictions would waste limited conservation resources on healthy or recovering coral.', bold_prefix='False alarm:')
add_bullet('Missing Dead coral (false negatives) could delay critical conservation responses.', bold_prefix='Missed cases:')
add_bullet('AI should help expert marine biologists, not replace them. Predictions should be reviewed by domain experts.', bold_prefix='Human oversight:')
add_bullet('Training on one geographic region could cause misclassification for coral from different oceans.', bold_prefix='Bias:')
add_bullet('Grad-CAM provides visual evidence for each prediction. Scientists can verify the decision basis before acting.', bold_prefix='Transparency:')
add_bullet('Low-confidence predictions (below 55%) should be sent for human review rather than automatic classification.', bold_prefix='Safety rule:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION N: CHAPTER 3 METHODOLOGY & ARCHITECTURE
# =====================================================
add_section_header('SECTION N: CHAPTER 3. METHODOLOGY & ARCHITECTURE (DEEP DIVE)')

# Q45
add_question(45, 'Explain your 10-stage methodology pipeline. Why did you design it this way?')
add_bullet('Phase 1 (Data Preparation): (1) Data Acquisition, (2) Data Preprocessing, (3) Data Splitting, (4) Data Augmentation.', bold_prefix='Phase 1:')
add_bullet('Phase 2 (Model Development): (5) Model Training with EfficientNetB0 + SWA, (6) Hyperparameter Tuning, (7) Model Export.', bold_prefix='Phase 2:')
add_bullet('Phase 3 (Analysis & Output): (8) Evaluation with Multi-Scale TTA, (9) Grad-CAM Interpretability, (10) Output Results & Deployment.', bold_prefix='Phase 3:')
add_bullet('Each stage feeds into the next. No stage can be skipped. The decision gate after tuning loops back to training if validation criteria are not met.', bold_prefix='Flow:')
add_bullet('This structure follows the standard ML workflow (prepare, train, evaluate) but adds coral-specific components: SWA, TTA, and Grad-CAM.', bold_prefix='Design reason:')
add_bullet('The flowchart (Figure 3.3) has explicit decision nodes. This makes the methodology iterative and self-correcting.', bold_prefix='Iterative:')

add_spacer()

# Q46
add_question(46, 'Describe your custom classification head architecture. Why this specific design?')
add_bullet('Structure: EfficientNetB0 backbone (ImageNet pretrained), then Global Average Pooling (GAP), then Dropout (0.4), then Dense(3, softmax).', bold_prefix='Architecture:')
add_bullet('GAP reduces the 7x7x1280 feature maps to a single 1280-dimensional vector. Better than Flatten because it avoids spatial overfitting.', bold_prefix='Why GAP:')
add_bullet('Dropout 0.4: randomly turns off 40% of neurons during training. Prevents co-adaptation and reduces overfitting on our small dataset.', bold_prefix='Why 0.4 dropout:')
add_bullet('Dense(3, softmax): 3 neurons for 3 classes. Softmax ensures outputs sum to 1.0, so they work as probabilities.', bold_prefix='Output layer:')
add_bullet('L2 regularisation (lambda = 0.0002): constrains weight magnitudes to prevent overfitting. Small value avoids over-penalising useful weights.', bold_prefix='L2:')
add_bullet('Top 100 layers unfrozen for fine-tuning. Higher layers adapt to coral features while lower layers keep their general visual knowledge.', bold_prefix='Fine-tuning:')
add_bullet('No extra Dense layers needed. Adding more layers increases parameters and overfitting risk. One Dense layer is enough for 3-class separation.', bold_prefix='Why minimal:')

add_spacer()

# Q47
add_question(47, 'Why did you unfreeze the final 100 layers? Why not all layers or fewer?')
add_bullet('All layers unfrozen: risks catastrophic forgetting. Pretrained low-level features (edges, textures) are universally useful and should be kept.', bold_prefix='Why not all:')
add_bullet('Too few layers (only the head): model cannot learn coral-specific features. The backbone stays entirely generic.', bold_prefix='Why not fewer:')
add_bullet('100 layers = the upper portion of EfficientNetB0. These contain higher-level, task-specific features that benefit most from fine-tuning.', bold_prefix='Why 100:')
add_bullet('Lower layers stay frozen. They retain universal visual primitives like colour gradients and edge detection from ImageNet.', bold_prefix='Frozen layers:')
add_bullet('This was tuned iteratively. 100 layers produced the best validation convergence during hyperparameter tuning.', bold_prefix='Empirically validated:')

add_spacer()

# Q48
add_question(48, 'Why Adam optimiser with learning rate 8x10^-5? Why not SGD or a higher LR?')
add_bullet('Adam adapts learning rates per-parameter and includes momentum. Converges faster than SGD on small datasets.', bold_prefix='Why Adam:')
add_bullet('SGD needs more manual tuning of LR and momentum scheduling. Adam is more forgiving and self-adapting.', bold_prefix='Why not SGD:')
add_bullet('LR = 8x10^-5: very small because we are fine-tuning pretrained weights. A large LR would destroy useful ImageNet features.', bold_prefix='Why small LR:')
add_bullet('Combined with Cosine Decay: LR smoothly decreases from 8x10^-5 toward zero over 30 epochs.', bold_prefix='LR schedule:')
add_bullet('This LR was selected during hyperparameter tuning. It produced stable convergence across all 5 seeds.', bold_prefix='Tested:')

add_spacer()

# Q49
add_question(49, 'Why 30 epochs? How did you determine this was sufficient?')
add_bullet('Training curves show convergence by epoch 22. Accuracy gap stabilises within +/-0.40 percentage points.', bold_prefix='Convergence:')
add_bullet('SWA collects weights from epochs 26 to 30 (final 5 epochs). These late-stage weights are near convergence.', bold_prefix='SWA needs this:')
add_bullet('Fewer epochs (e.g., 15): model not fully converged. SWA would average premature weights.', bold_prefix='Why not fewer:')
add_bullet('More epochs (e.g., 50, 100): unnecessary. Cosine decay has already reduced LR near zero by epoch 30.', bold_prefix='Why not more:')
add_bullet('Validation loss showed no upward divergence at epoch 30. Confirms no overfitting and no need for early stopping.', bold_prefix='No overfitting:')

add_spacer()

# Q50
add_question(50, 'Why label smoothing epsilon = 0.05? What does it do?')
add_bullet('Replaces hard labels [1, 0, 0] with soft labels [0.95, 0.025, 0.025]. Prevents the model from becoming overconfident.', bold_prefix='What it does:')
add_bullet('Without smoothing: model is trained to output 100% confidence. This leads to sharp distributions that don\'t generalise well.', bold_prefix='Problem:')
add_bullet('epsilon = 0.05 is a mild regularisation. Subtracts 5% from the true class and spreads it to the others.', bold_prefix='Value:')
add_bullet('Higher values like 0.2 would over-smooth and reduce the model\'s ability to discriminate between classes.', bold_prefix='Why not higher:')
add_bullet('Particularly useful for the Bleached-Dead boundary where visual ambiguity exists. Encourages expressing uncertainty.', bold_prefix='Relevant here:')

add_spacer()

# Q51
add_question(51, 'Why bilinear interpolation for resizing? Why not bicubic or nearest-neighbour?')
add_bullet('Bilinear: averages 4 nearest pixels. Smooth and computationally efficient.', bold_prefix='Bilinear:')
add_bullet('Nearest-neighbour: copies pixels without averaging. Produces blocky, pixelated results with aliasing.', bold_prefix='Why not NN:')
add_bullet('Bicubic: uses 16 pixels. Smoother but more expensive and can introduce ringing artefacts at edges.', bold_prefix='Why not bicubic:')
add_bullet('Bilinear is the standard for EfficientNet and most CNNs. Consistent with how ImageNet weights were trained.', bold_prefix='Standard:')
add_bullet('Coral textures benefit from smooth resizing. Preserves colour gradients important for health classification.', bold_prefix='Relevance:')

add_spacer()

# Q52
add_question(52, 'Why BGR to RGB conversion? Why not just load in RGB directly?')
add_bullet('OpenCV (cv2.imread) loads images in BGR format by default. Historical convention from old camera hardware.', bold_prefix='Why BGR:')
add_bullet('EfficientNetB0 expects RGB input. All ImageNet pretrained weights were learned on RGB images.', bold_prefix='Model expects:')
add_bullet('If you feed BGR to an RGB model, the blue and red channels are swapped. The model sees wrong colours.', bold_prefix='Consequence:')
add_bullet('Colour is critical for coral: healthy = vibrant, bleached = white/pale. Swapped channels would destroy these cues.', bold_prefix='Impact:')
add_bullet('Simple fix: cv2.cvtColor(image, cv2.COLOR_BGR2RGB). Standard practice in all OpenCV + deep learning pipelines.', bold_prefix='Fix:')

add_spacer()

# Q53
add_question(53, 'Explain the Grad-CAM mathematical formulation. How does it generate heatmaps?')
add_bullet('Step 1 (Forward pass): compute the class score yc for target class c.', bold_prefix='Forward:')
add_bullet('Step 2 (Backward pass): compute gradients of yc with respect to each feature map Ak at every spatial position.', bold_prefix='Backward:')
add_bullet('Step 3 (Weight computation): global average pool the gradients to get channel-wise importance weights alpha_k.', bold_prefix='Weights:')
add_bullet('Step 4 (Heatmap): weighted combination of feature maps. Apply ReLU to keep only positive contributions. L = ReLU(sum of alpha_k x Ak).', bold_prefix='Combine:')
add_bullet('Step 5 (Upsample): bilinear interpolation from 7x7 (feature map size) to 224x224 (input size).', bold_prefix='Resize:')
add_bullet('Step 6 (Overlay): apply JET colourmap (blue=low, red=high activation). Alpha-blend onto original image.', bold_prefix='Display:')
add_bullet('ReLU is critical: without it, regions that suppress the class would appear as activations, creating misleading heatmaps.', bold_prefix='Why ReLU:')

add_spacer()

# Q54
add_question(54, 'Why did you target the "top_conv" layer for Grad-CAM? Why not an earlier layer?')
add_bullet('top_conv is EfficientNetB0\'s final convolutional layer. It has the most semantically rich feature representations.', bold_prefix='Choice:')
add_bullet('Earlier layers capture low-level features (edges, textures). Grad-CAM on these would show generic patterns, not class-discriminative regions.', bold_prefix='Why not earlier:')
add_bullet('Deeper layers capture high-level, class-specific features. These are what determine the classification decision.', bold_prefix='Why deepest:')
add_bullet('Selvaraju et al. (2020) recommend the last conv layer as the optimal Grad-CAM target for all CNN architectures.', bold_prefix='Literature:')
add_bullet('top_conv feature maps have 7x7 spatial resolution. Enough for meaningful localisation while being efficient.', bold_prefix='Resolution:')

add_spacer()

# Q55
add_question(55, 'What are the expected Grad-CAM focal regions for each class and why?')
add_bullet('Healthy: pigmented coral structures (branches, fans). Zooxanthellae give healthy coral its vibrant colour.', bold_prefix='Healthy:')
add_bullet('Bleached: broad activation across pale/white surface. Bleaching causes colony-wide pigmentation loss, so the attention should be spread out.', bold_prefix='Bleached:')
add_bullet('Dead: exposed calcium carbonate skeleton, often with turf algae. Structural texture without living tissue colour.', bold_prefix='Dead:')
add_bullet('These expectations come from established coral biology. Table 3.9 in the report maps each class to its biological indicators.', bold_prefix='Basis:')
add_bullet('Our ensemble Grad-CAM results matched all three expected patterns. Confirms the model learned biologically meaningful features.', bold_prefix='Confirmed:')

add_spacer()

# Q56
add_question(56, 'How does Multi-Scale TTA handle the 256x256 resolution? The model expects 224x224.')
add_bullet('Image is first resized to 256x256, then centre-cropped to 224x224. This maintains the model\'s expected input size.', bold_prefix='Process:')
add_bullet('Centre-cropping at 256 to 224 keeps 76.6% of the spatial area. The model sees a slightly "zoomed in" view of the centre.', bold_prefix='Crop ratio:')
add_bullet('Two perspectives: 224x224 native sees the full image. 256x256 centre-crop focuses on the central coral subject.', bold_prefix='Two views:')
add_bullet('Combined with horizontal flip: 2 scales x 2 orientations = 4 views per image per model.', bold_prefix='Total:')

add_spacer()

# Q57
add_question(57, 'Why did you use Keras ImageDataGenerator and not albumentations or torchvision?')
add_bullet('ImageDataGenerator is built into the TensorFlow/Keras pipeline. No additional libraries needed.', bold_prefix='Built-in:')
add_bullet('On-the-fly augmentation: applies transforms during training, not before. Saves disk space and gives different augments each epoch.', bold_prefix='Efficient:')
add_bullet('Albumentations has more advanced transforms but requires extra installation and manual batching logic.', bold_prefix='Why not albumentations:')
add_bullet('Torchvision needs PyTorch. Our entire stack is TensorFlow/Keras. Mixing frameworks causes compatibility issues.', bold_prefix='Why not torchvision:')
add_bullet('For our needs (rotation, flip, zoom, brightness), ImageDataGenerator provides everything required.', bold_prefix='Sufficient:')

add_spacer()

# Q58
add_question(58, 'Why TensorFlow/Keras instead of PyTorch?')
add_bullet('Keras has a high-level API: model.fit(), callbacks, ImageDataGenerator. Rapid prototyping with less code.', bold_prefix='Easy:')
add_bullet('EfficientNetB0 pretrained weights available via tf.keras.applications. No manual weight downloading needed.', bold_prefix='Model access:')
add_bullet('ModelCheckpoint and LearningRateScheduler callbacks are built-in. No custom training loop required.', bold_prefix='Callbacks:')
add_bullet('HDF5 model export is compatible with Flask deployment. Seamless training-to-deployment pipeline.', bold_prefix='Deployment:')
add_bullet('PyTorch has more flexibility but needs manual training loops and more boilerplate code. Overkill for our scope.', bold_prefix='PyTorch:')

add_spacer()

# Q59
add_question(59, 'Explain the SWA implementation. When exactly are weights collected?')
add_bullet('SWA is implemented via a custom Keras callback that runs at the end of each epoch.', bold_prefix='Implementation:')
add_bullet('Weights collected from epochs 26 to 30 (final 5 of 30 epochs). Model has converged by this point.', bold_prefix='When:')
add_bullet('At each collection epoch: current weights are added to a running arithmetic mean of previously collected weights.', bold_prefix='Averaging:')
add_bullet('After epoch 30: the SWA-averaged weights are saved as the final checkpoint. This represents the centroid of 5 late-stage weight snapshots.', bold_prefix='Final:')
add_bullet('This centroid sits in a flatter region of the loss landscape. Flatter = better generalisation (Izmailov et al., 2018).', bold_prefix='Theory:')
add_bullet('This process repeats independently for each of the 5 seeds (42-46), producing 5 distinct SWA checkpoints.', bold_prefix='Per seed:')

add_spacer()

# Q60
add_question(60, 'What is your hyperparameter tuning process? Is it grid search or manual?')
add_bullet('Manual iterative tuning. Not grid search or Bayesian optimisation.', bold_prefix='Method:')
add_bullet('Parameters tuned: dropout rate, L2 coefficient, label smoothing, initial LR, unfrozen layers, augmentation intensity, oversampling rates, class weights.', bold_prefix='Parameters:')
add_bullet('Process: train, inspect validation curves, adjust parameters, retrain, repeat until convergence criteria met.', bold_prefix='Iterative:')
add_bullet('Acceptance gate: validation accuracy, loss, AND macro F1 must all meet satisfactory thresholds.', bold_prefix='Criteria:')
add_bullet('Grid search is impractical: 8 hyperparameters + 5-seed training = too many combinations to try exhaustively.', bold_prefix='Why not grid:')
add_bullet('Table 3.7 in the report shows baseline vs final values for each hyperparameter with reasons for changes.', bold_prefix='Documented:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION O: CHAPTER 4 RESULTS & DISCUSSION
# =====================================================
add_section_header('SECTION O: CHAPTER 4. RESULTS & DISCUSSION (DEEP DIVE)')

# Q61
add_question(61, 'Your training accuracy is LOWER than validation accuracy. Isn\'t that a problem?')
add_bullet('No. This is the expected result of our high-regularisation training strategy.', bold_prefix='Short answer:')
add_bullet('Cause 1: Hard-Example Oversampling at 30x makes training batches artificially harder.', bold_prefix='Oversampling:')
add_bullet('Cause 2: Augmentation (rotation, flip, zoom, brightness) adds noise to training images only.', bold_prefix='Augmentation:')
add_bullet('Cause 3: Dropout 0.4 disables 40% of neurons during training. At validation, full network capacity is available.', bold_prefix='Dropout:')
add_bullet('Validation set: no augmentation, no oversampling, no dropout. It sees clean original images with full network.', bold_prefix='Validation is easier:')
add_bullet('Initial gap: -19.31 percentage points at epoch 1. Stabilised to +/-0.40 by epoch 22. Shows progressive convergence.', bold_prefix='Numbers:')
add_bullet('If this were underfitting, both metrics would be poor. But validation accuracy exceeds 98%.', bold_prefix='Not underfitting:')

add_spacer()

# Q62
add_question(62, 'Explain the accuracy gap progression. What does the -19.31 point gap at epoch 1 mean?')
add_bullet('Epoch 1: training accuracy ~75%, validation accuracy ~94%. Gap = -19.31 percentage points (validation higher).', bold_prefix='Epoch 1:')
add_bullet('This large gap reflects the regularisation burden at training start. Model has not adapted to the hard augmented data yet.', bold_prefix='Explanation:')
add_bullet('As training continues: model learns to handle augmented inputs. Training accuracy rises. Gap narrows.', bold_prefix='Progression:')
add_bullet('By epoch 22: gap stabilises to +/-0.40 pp. Model has fully adapted to the training distribution.', bold_prefix='Convergence:')
add_bullet('The narrowing gap proves the model is learning robust features, not just memorising easy patterns.', bold_prefix='Meaning:')
add_bullet('Loss gap also narrowed: from 0.354 at epoch 1 to 0.10 at epoch 30.', bold_prefix='Loss confirms:')

add_spacer()

# Q63
add_question(63, 'Your Dead class recall is 93.33%. Why isn\'t it 100%? What went wrong?')
add_bullet('1 Dead sample out of 15 was predicted as Bleached. Model gave 51.3% confidence (very uncertain).', bold_prefix='What happened:')
add_bullet('Root cause: genuine visual ambiguity. Dead and Bleached coral share reduced pigmentation and structural degradation.', bold_prefix='Visual overlap:')
add_bullet('Grad-CAM confirmed the model looked at the right region (coral body) but chose the wrong class.', bold_prefix='Not attention failure:')
add_bullet('51.3% confidence is near the 3-class chance level (33.3%). The 5 seeds did not agree on a class.', bold_prefix='Low confidence:')
add_bullet('Biologically expected: the Bleached to Dead transition is gradual. Some images capture intermediate states.', bold_prefix='Biology:')
add_bullet('Fix for deployment: route predictions below 55% confidence to human expert review.', bold_prefix='Safeguard:')

add_spacer()

# Q64
add_question(64, 'Dead precision is 1.000. What does this operationally mean?')
add_bullet('Zero false Dead predictions. No Healthy or Bleached coral was ever wrongly labelled as Dead across all 159 test images.', bold_prefix='Meaning:')
add_bullet('If the model says "Dead", it is always correct. Conservation teams can trust Dead predictions fully.', bold_prefix='Trust:')
add_bullet('False Dead predictions would waste limited resources on healthy/recovering coral. Our model avoids this completely.', bold_prefix='Resource impact:')
add_bullet('Trade-off: the model occasionally misses Dead coral (1/15) but never falsely triggers a Dead alarm.', bold_prefix='Trade-off:')
add_bullet('In conservation, this is preferable. False alarms erode trust in the system over time.', bold_prefix='Better:')

add_spacer()

# Q65
add_question(65, 'Explain the 13.20 percentage point improvement from baseline. What contributed most?')
add_bullet('Baseline: single EfficientNetB0, standard setup, no SWA, no oversampling, no TTA. Result: 84.91% accuracy, 24 errors.', bold_prefix='Baseline:')
add_bullet('Final: 5-seed SWA ensemble, 30x Dead oversampling, cosine decay, class weights, Multi-Scale TTA. Result: 98.11%, 3 errors.', bold_prefix='Final:')
add_bullet('Biggest contributor: Ensemble + SWA. Ablation shows 1 seed = 96.23% (6 errors) vs 5 seeds = 98.11% (3 errors). Accounts for ~11 pp gain.', bold_prefix='#1 Ensemble:')
add_bullet('#2: Hard-Example Oversampling addressed the Dead class imbalance. Baseline likely failed on most Dead samples.', bold_prefix='#2 Oversampling:')
add_bullet('#3: Cosine Decay enabled smoother convergence vs baseline\'s LR strategy.', bold_prefix='#3 LR schedule:')
add_bullet('Macro F1 improved from 79.19% to 97.69%. Balanced improvement across all 3 classes.', bold_prefix='Balanced:')

add_spacer()

# Q66
add_question(66, 'Why do all 3 misclassifications have sub-55% confidence? What does this tell us?')
add_bullet('Ensemble confidence reflects agreement among 5 seeds. High confidence = consensus. Low confidence = disagreement.', bold_prefix='How it works:')
add_bullet('Bleached-as-Healthy errors: 54.3% confidence. Seeds were split between Healthy and Bleached.', bold_prefix='Error 1&2:')
add_bullet('Dead-as-Bleached error: 51.3% confidence. Near the 3-class chance threshold (33.3%). No consensus.', bold_prefix='Error 3:')
add_bullet('The ensemble "knows when it doesn\'t know." Low confidence reliably indicates uncertain predictions.', bold_prefix='Self-aware:')
add_bullet('Deployment rule: if confidence < 55%, send prediction to human expert review instead of auto-classifying.', bold_prefix='Confidence gate:')
add_bullet('All 156 correct predictions had high confidence (97%+). Clear separation between confident-correct and uncertain-wrong.', bold_prefix='Separation:')

add_spacer()

# Q67
add_question(67, 'How does the confusion matrix error pattern relate to coral biology?')
add_bullet('Errors only flow between adjacent states: Bleached to Healthy (2 cases), Dead to Bleached (1 case).', bold_prefix='Directional:')
add_bullet('No "skip" errors: no Dead to Healthy or Healthy to Dead misclassifications.', bold_prefix='No skips:')
add_bullet('Biology: coral health is a continuum. Healthy to Bleached to Dead is a progressive degradation pathway.', bold_prefix='Continuum:')
add_bullet('Adjacent states share overlapping visual features. Early bleaching looks like faded healthy coral. Late bleaching looks like early death.', bold_prefix='Feature overlap:')
add_bullet('The model only struggles where visual distinctions are genuinely ambiguous. Mirrors biological reality.', bold_prefix='Validates model:')
add_bullet('This confirms the model learned the biological progression order, not arbitrary class boundaries.', bold_prefix='Learned correctly:')

add_spacer()

# Q68
add_question(68, 'In the discussion, you mention the baseline Healthy Grad-CAM focused on a clownfish. Elaborate on this finding.')
add_bullet('Baseline predicted Healthy at 75.2% confidence (correct class). But Grad-CAM showed it focused on a clownfish, not coral.', bold_prefix='What happened:')
add_bullet('The coral body received almost no attention. The model was classifying based on the presence of fish.', bold_prefix='Wrong feature:')
add_bullet('Spurious correlation: clownfish appear with healthy coral in training data. Model learned "clownfish = healthy" instead of "pigmented tissue = healthy."', bold_prefix='Root cause:')
add_bullet('In deployment: a healthy coral image without fish would likely be misclassified. This shortcut does not generalise.', bold_prefix='Deployment failure:')
add_bullet('This is the strongest argument for why Objective 2 (Grad-CAM) is necessary. Accuracy alone would mark this as a "success."', bold_prefix='Key point:')
add_bullet('Our ensemble fixed this: 5-seed SWA + oversampling forced the model to learn coral tissue features rather than scene-level shortcuts.', bold_prefix='Fixed:')
add_bullet('Samek et al. (2021): XAI reveals reliance on spurious correlations that accuracy metrics cannot detect.', bold_prefix='Literature:')

add_spacer()

# Q69
add_question(69, 'How do you address the statistical uncertainty from only 159 test samples?')
add_bullet('Acknowledged: 159 test images is small. Dead class has only 15 samples. Each error shifts recall by 6.67 pp.', bold_prefix='Limitation:')
add_bullet('But the stratified split preserves class proportions. Test set is representative of the full dataset distribution.', bold_prefix='Representative:')
add_bullet('Multiple converging evidence sources: (1) ablation study across 8 configs, (2) 5-seed consistency, (3) Grad-CAM biological validation.', bold_prefix='Multiple sources:')
add_bullet('The error pattern (directional, adjacent-state only) is biologically coherent. Unlikely to be a statistical artefact.', bold_prefix='Pattern valid:')
add_bullet('Future work: cross-regional validation with larger, diverse test sets for broader statistical confidence.', bold_prefix='Future fix:')

add_spacer()

# Q70
add_question(70, 'Summarise how your results satisfy each of the three research objectives.')
add_bullet('Objective 1 (model): EfficientNetB0 + 5-seed SWA ensemble. Result: 98.11% accuracy, +13.20 pp over baseline, Macro F1 = 0.9769. ACHIEVED.', bold_prefix='Obj 1:')
add_bullet('Objective 2 (Grad-CAM): biologically coherent heatmaps. Pigmented tissue (Healthy), pale surface (Bleached), skeletal formations (Dead). Baseline comparison proved ensemble learns correct features. ACHIEVED.', bold_prefix='Obj 2:')
add_bullet('Objective 3 (multi-metric evaluation): accuracy, precision, recall, F1, Macro F1, weighted F1, confusion matrix all reported per-class. Dead precision = 1.000, all class F1 > 0.965. ACHIEVED.', bold_prefix='Obj 3:')
add_bullet('Bonus Obj 4 (deployment): Flask web app with real-time inference, confidence scoring, and Grad-CAM overlay. ACHIEVED.', bold_prefix='Obj 4:')

add_spacer()

    # doc.add_page_break()

# =====================================================
# SECTION P: QUICK REFERENCE STATISTICS
# =====================================================
add_section_header('SECTION P: QUICK REFERENCE. KEY STATISTICS')

stats = [
    ('Test Accuracy', '98.11% (156/159 correct)'),
    ('Baseline Accuracy', '84.91% (24 errors)'),
    ('Accuracy Improvement', '+13.20 percentage points'),
    ('Macro F1-Score', '0.9769'),
    ('Weighted F1-Score', '0.9810'),
    ('Healthy F1', '0.9863 (Precision: 0.973, Recall: 1.000)'),
    ('Bleached F1', '0.9790 (Precision: 0.986, Recall: 0.972)'),
    ('Dead F1', '0.9655 (Precision: 1.000, Recall: 0.933)'),
    ('Total Misclassifications', '3 (2 Bleached>Healthy, 1 Dead>Bleached)'),
    ('Dead False Positives', '0 (zero false Dead predictions)'),
    ('Ensemble Size', '5 seeds (42, 43, 44, 45, 46)'),
    ('Epochs per Seed', '30'),
    ('SWA Window', 'Epochs 26-30 (final 5 epochs)'),
    ('Training Time per Seed', '~12 minutes'),
    ('Total Training Time', '~60 minutes (5 seeds)'),
    ('TTA Views per Image', '4 (224+256 x original+flip)'),
    ('Total Predictions per Image', '20 (5 seeds x 4 TTA views)'),
    ('Mean Inference Time', '10.38 ms/image'),
    ('Input Resolution', '224 x 224 x 3 RGB'),
    ('EfficientNetB0 Parameters', '~5.3M'),
    ('Unfrozen Layers', 'Final 100 layers'),
    ('Dropout Rate', '0.4'),
    ('L2 Regularisation', 'lambda = 0.0002'),
    ('Label Smoothing', 'epsilon = 0.05'),
    ('Optimiser', 'Adam (LR = 8x10^-5)'),
    ('LR Schedule', 'Cosine Decay over 30 epochs'),
    ('Batch Size', '16'),
    ('Loss Function', 'Categorical Cross-Entropy'),
    ('Dataset Total', '1,582 images'),
    ('Dataset Split', '1,265 train / 158 val / 159 test'),
    ('Class Distribution', 'Healthy: 712, Bleached: 720, Dead: 150'),
    ('Dead Oversampling Rate', '30x'),
    ('Healthy/Bleached Hard-Example Rate', '20x'),
    ('Dead Class Weight Multiplier', '1.3x'),
    ('Grad-CAM Target Layer', 'top_conv (final conv layer)'),
    ('Grad-CAM Colourmap', 'JET (blue>green>yellow>red)'),
    ('Ensemble Confidence (Correct)', '96.6%-99.4%'),
    ('Baseline Confidence', '45.5%-75.2%'),
    ('Misclassification Confidence', '<55% (all 3 errors)'),
    ('GPU', 'NVIDIA RTX 3070 (8 GB VRAM)'),
    ('CUDA / cuDNN', '11.8 / 8.6'),
    ('Framework', 'TensorFlow 2.12 / Keras'),
    ('Python', '3.10'),
    ('Deployment', 'Flask + Flask-CORS'),
]

table = doc.add_table(rows=1, cols=2, style='Table Grid')
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Metric / Item'
hdr_cells[1].text = 'Value'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for item, value in stats:
    row_cells = table.add_row().cells
    row_cells[0].text = item
    row_cells[1].text = value

# Save
output_path = r'C:\Users\ZeeqRyz\Desktop\CHI\BASEPROJECT\Fyp_Report\Deliverables\Viva_QA_Preparation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Total questions: 70')
