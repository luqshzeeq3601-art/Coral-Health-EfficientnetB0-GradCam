from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Coral_Poster_Presentation_Script_UPEX_2026.docx"


SECTIONS = [
    {
        "heading": "Opening",
        "bullets": [
            "Good morning judges.",
            "My name is Muhammad Luqman Haziq Bin Muhamad Lofi.",
            "My project is Coral Reef Health Assessment via Convolutional Neural Network-Based Image Analysis.",
            "This project uses AI to classify coral reef images and explain the result using Grad-CAM heatmaps.",
        ],
    },
    {
        "heading": "Problem Statement",
        "bullets": [
            "Coral monitoring is still mostly manual.",
            "It needs expert divers, field surveys, time, and cost.",
            "Coral bleaching is a serious threat caused by climate change and environmental stress.",
            "Many AI systems also have a black-box problem, where they give a result without explaining why.",
        ],
    },
    {
        "heading": "Objectives",
        "bullets": [
            "Classify coral images into Healthy, Bleached, and Dead.",
            "Provide explainable AI using Grad-CAM.",
            "Evaluate performance using accuracy, precision, recall, F1-score, and confusion matrix.",
        ],
    },
    {
        "heading": "Methodology",
        "bullets": [
            "Step 1: collect and organize coral images into Healthy, Bleached, and Dead classes.",
            "Step 2: split the dataset into training, validation, and test sets.",
            "Step 3: preprocess images by resizing to 224 x 224 and converting BGR to RGB.",
            "Step 4: apply augmentation such as flipping, rotation, zoom, shift, and brightness adjustment.",
            "Step 5: handle hard examples with oversampling and class weighting.",
            "Step 6: train an EfficientNetB0 ensemble with five random seeds.",
            "Step 7: apply SWA to improve generalization.",
            "Step 8: evaluate with standard metrics and Grad-CAM explainability.",
        ],
    },
    {
        "heading": "Results",
        "bullets": [
            "The user uploads a coral image.",
            "The system predicts the class and shows the confidence score.",
            "Grad-CAM highlights the important area used by the model.",
            "ReefGuide support helps explain the result in simple language.",
            "The web application can classify Healthy, Bleached, and Dead coral images.",
            "Healthy coral is supported by natural coral texture.",
            "Bleached coral is supported by pale coral regions.",
            "Dead coral is supported by degraded coral areas.",
            "The system provides prediction, confidence, class probabilities, and Grad-CAM explanation for each uploaded image.",
        ],
    },
    {
        "heading": "Commercialization and Impact",
        "bullets": [
            "The main product is a Coral Health AI Platform.",
            "It supports automated assessment, explainable AI, web deployment, and scalable monitoring.",
            "Target users include marine agencies, environmental monitoring teams, research institutions, and coral restoration groups.",
            "Use cases include coral health assessment, bleaching detection, and reef monitoring.",
            "The project contributes to society, community, industry, and environmental sustainability.",
        ],
    },
    {
        "heading": "Closing",
        "bullets": [
            "This project provides an AI-based coral health assessment system that is predictive and explainable.",
            "It helps users classify coral condition, understand the AI decision, and support faster reef monitoring.",
            "Thank you.",
        ],
    },
]


SHORT_BULLETS = [
    "Good morning judges.",
    "My project is Coral Reef Health Assessment via Convolutional Neural Network-Based Image Analysis.",
    "The problem is that coral monitoring is manual, expensive, and time-consuming, while many AI systems are difficult to trust because they work like a black box.",
    "My system classifies coral images into Healthy, Bleached, and Dead, and explains the prediction using Grad-CAM heatmaps.",
    "The pipeline covers data collection, splitting, preprocessing, augmentation, class balancing, EfficientNetB0 ensemble training, SWA optimization, final evaluation, and Grad-CAM explainability.",
    "In the web application, users upload a coral image, receive a prediction with confidence and class probabilities, and can view the Grad-CAM focus area.",
    "ReefGuide support helps explain the result in simple language.",
    "The platform can support marine agencies, environmental monitoring teams, research institutions, and coral restoration groups.",
    "Overall, this project provides a faster and more explainable AI tool for coral reef health monitoring.",
    "Thank you.",
]


KEY_TERMS = [
    ("CNN", "A deep learning model used for image classification."),
    ("EfficientNetB0", "The lightweight CNN backbone used for coral classification."),
    ("Grad-CAM", "A heatmap method that shows which image regions influenced the prediction."),
    ("SWA", "Stochastic Weight Averaging, used to improve model generalization."),
    ("Ensemble", "Five models trained with different seeds and averaged for more stable prediction."),
    ("TTA", "Test-Time Augmentation, where multiple image views are predicted and averaged."),
    ("Class Weighting", "Higher penalty for difficult or minority classes during training."),
]


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)

    for name, size in [("Heading 1", 15), ("Heading 2", 12)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)

    if "Bullet Text" not in doc.styles:
        bullet = doc.styles.add_style("Bullet Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = doc.styles["Bullet Text"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.space_after = Pt(1)
    bullet.paragraph_format.left_indent = Inches(0.25)
    bullet.paragraph_format.first_line_indent = Inches(-0.12)


def add_bullets(doc, bullets):
    for bullet in bullets:
        p = doc.add_paragraph(style="Bullet Text")
        p.style = "List Bullet"
        p.add_run(bullet)


doc = Document()
setup(doc)

doc.add_paragraph("Coral Poster Presentation Script", style="Title")
doc.add_paragraph("UPEX 2026 | Bullet-point script for quick review", style="Subtitle")

intro = doc.add_paragraph()
intro.add_run("How to use this file: ").bold = True
intro.add_run("Practice the full bullet script for a short presentation. Use the short version if you are tight on time.")

doc.add_heading("Full Poster Script", level=1)
for section in SECTIONS:
    doc.add_heading(section["heading"], level=2)
    add_bullets(doc, section["bullets"])

doc.add_page_break()
doc.add_heading("Short Version If Time Is Limited", level=1)
add_bullets(doc, SHORT_BULLETS)

doc.add_heading("Poster Walkthrough Order", level=1)
walkthrough = [
    "Start at the title and introduce the project.",
    "Move to Problem Statement: manual monitoring, bleaching threat, AI black-box problem.",
    "Move to Objectives: classify, explain, evaluate.",
    "Move to Methodology: explain the 8-step pipeline.",
    "Move to Results: show upload flow, prediction, confidence, Grad-CAM, and ReefGuide.",
    "End at Commercialization and Impact: product, target users, use cases, and contribution.",
]
for item in walkthrough:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("Key Terms to Review", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Term"
table.rows[0].cells[1].text = "Simple Meaning"
for cell in table.rows[0].cells:
    shade(cell, "E8EEF5")
for term, meaning in KEY_TERMS:
    row = table.add_row().cells
    row[0].text = term
    row[1].text = meaning
table_geometry(table, [2200, 6800])

doc.add_heading("Final Reminder", level=1)
final_points = [
    "Do not read too fast.",
    "Use simple words: classify, explain, evaluate, deploy.",
    "If judges ask about accuracy, mention that strong test performance still needs wider real-world validation.",
    "If judges ask about trust, point to Grad-CAM and confidence score.",
    "If judges ask about usefulness, explain marine agencies, monitoring teams, research institutions, and coral restoration groups.",
]
for point in final_points:
    doc.add_paragraph(point, style="List Bullet")

footer = doc.sections[0].footer.paragraphs[0]
footer.text = "UPEX 2026 - Coral Reef Health Assessment Poster Script"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print(OUT)
